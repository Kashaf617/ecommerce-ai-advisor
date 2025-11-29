"""
Helper utilities for the E-Commerce Business Automation Platform
"""
import csv
import json
from pathlib import Path
from typing import List, Dict, Any
from datetime import datetime
import pandas as pd

def save_to_csv(data: List[Dict[str, Any]], filepath: Path, mode: str = 'w') -> bool:
    """
    Save data to CSV file
    
    Args:
        data: List of dictionaries to save
        filepath: Path to CSV file
        mode: File mode ('w' for write, 'a' for append)
    
    Returns:
        True if successful, False otherwise
    """
    if not data:
        return False
    
    try:
        filepath.parent.mkdir(parents=True, exist_ok=True)
        
        with open(filepath, mode, newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=data[0].keys())
            if mode == 'w' or not filepath.exists() or filepath.stat().st_size == 0:
                writer.writeheader()
            writer.writerows(data)
        return True
    except Exception as e:
        print(f"Error saving to CSV: {e}")
        return False

def load_from_csv(filepath: Path) -> List[Dict[str, Any]]:
    """
    Load data from CSV file
    
    Args:
        filepath: Path to CSV file
    
    Returns:
        List of dictionaries
    """
    if not filepath.exists():
        return []
    
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return list(csv.DictReader(f))
    except Exception as e:
        print(f"Error loading from CSV: {e}")
        return []

def save_to_json(data: Any, filepath: Path) -> bool:
    """Save data to JSON file"""
    try:
        filepath.parent.mkdir(parents=True, exist_ok=True)
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Error saving to JSON: {e}")
        return False

def load_from_json(filepath: Path) -> Any:
    """Load data from JSON file"""
    if not filepath.exists():
        return None
    
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"Error loading from JSON: {e}")
        return None

def clean_text(text: str) -> str:
    """Clean and normalize text"""
    if not text:
        return ""
    return ' '.join(text.strip().split())

def format_currency(amount: float, currency: str = 'USD') -> str:
    """Format amount as currency"""
    symbols = {'USD': '$', 'PKR': 'Rs.', 'EUR': '€', 'GBP': '£'}
    symbol = symbols.get(currency, currency)
    return f"{symbol}{amount:,.2f}"

def calculate_percentage_change(old_value: float, new_value: float) -> float:
    """Calculate percentage change between two values"""
    if old_value == 0:
        return 0
    return ((new_value - old_value) / old_value) * 100

def get_timestamp() -> str:
    """Get current timestamp as string"""
    return datetime.now().strftime('%Y-%m-%d %H:%M:%S')

def merge_csv_files(input_files: List[Path], output_file: Path) -> bool:
    """Merge multiple CSV files into one"""
    try:
        dfs = []
        for file in input_files:
            if file.exists():
                df = pd.read_csv(file)
                dfs.append(df)
        
        if dfs:
            merged_df = pd.concat(dfs, ignore_index=True)
            output_file.parent.mkdir(parents=True, exist_ok=True)
            merged_df.to_csv(output_file, index=False)
            return True
        return False
    except Exception as e:
        print(f"Error merging CSV files: {e}")
        return False

def export_to_docx(results: Dict[str, Any], filepath: Path) -> bool:
    """
    Export analysis results to a formatted Word document
    
    Args:
        results: Analysis results dictionary
        filepath: Path to save DOCX file
    
    Returns:
        True if successful, False otherwise
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading('E-Commerce Business Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Input summary
        doc.add_heading('Product Information', level=1)
        input_data = results.get('input', {})
        
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Light Grid Accent 1'
        
        input_items = [
            ('Product Name', input_data.get('product_name', 'N/A')),
            ('Category', input_data.get('category', 'N/A')),
            ('Price', f"{input_data.get('currency_symbol', '$')}{input_data.get('price', 0):,.2f}"),
            ('Budget', f"{input_data.get('currency_symbol', '$')}{input_data.get('budget', 0):,.2f}"),
            ('Quantity', str(input_data.get('quantity', 0))),
            ('Target Market', input_data.get('target_market', 'N/A').title()),
            ('Analysis Date', results.get('timestamp', 'N/A'))
        ]
        
        for i, (key, value) in enumerate(input_items):
            table.rows[i].cells[0].text = key
            table.rows[i].cells[1].text = str(value)
        
        doc.add_paragraph()
        
        # Module 2: Trends
        if 'module_2_trends' in results:
            doc.add_heading('Market Trends & Demand Forecast', level=1)
            trends = results['module_2_trends'].get('category_trends', {})
            forecast = results['module_2_trends'].get('demand_forecast', {})
            currency_symbol = input_data.get('currency_symbol', '$')
            
            p = doc.add_paragraph()
            p.add_run(f"Average Market Price: ").bold = True
            p.add_run(f"{currency_symbol}{trends.get('average_price', 0):,.2f}\n")
            p.add_run(f"Average Rating: ").bold = True
            p.add_run(f"{trends.get('average_rating', 0):.1f}/5.0\n")
            p.add_run(f"Demand Level: ").bold = True
            p.add_run(f"{forecast.get('demand_level', 'N/A')}\n")
            
            if forecast.get('recommendations'):
                doc.add_heading('Recommendations:', level=2)
                for rec in forecast['recommendations'][:5]:
                    doc.add_paragraph(rec, style='List Bullet')
        
        # Module 3: Competitors
        if 'module_3_competitors' in results:
            doc.add_heading('Competitor Analysis', level=1)
            comp = results['module_3_competitors']
            analysis = comp.get('competitive_analysis', {})
            
            p = doc.add_paragraph()
            p.add_run(f"Total Competitors: ").bold = True
            p.add_run(f"{comp['total_competitors']}\n")
            p.add_run(f"Your Price Positioning: ").bold = True
            p.add_run(f"{analysis.get('price_positioning', 'N/A')}\n")
            
            if comp.get('top_competitors'):
                doc.add_heading('Top Competitors:', level=2)
                comp_table = doc.add_table(rows=1, cols=4)
                comp_table.style = 'Light Grid Accent 1'
                
                hdr_cells = comp_table.rows[0].cells
                hdr_cells[0].text = 'Seller'
                hdr_cells[1].text = 'Score'
                hdr_cells[2].text = 'Rating'
                hdr_cells[3].text = 'Products'
                
                for competitor in comp['top_competitors'][:5]:
                    row_cells = comp_table.add_row().cells
                    row_cells[0].text = competitor['seller_name']
                    row_cells[1].text = f"{competitor['competitive_score']:.1f}"
                    row_cells[2].text = f"{competitor['average_rating']:.1f}"
                    row_cells[3].text = str(competitor['product_count'])
        
        # Module 4: Suppliers
        if 'module_4_suppliers' in results:
            doc.add_heading('Recommended Suppliers', level=1)
            suppliers = results['module_4_suppliers'].get('recommended_suppliers', [])
            currency_symbol = input_data.get('currency_symbol', '$')
            
            if suppliers:
                for i, supplier in enumerate(suppliers[:3], 1):
                    doc.add_heading(f"{i}. {supplier['name']}", level=2)
                    p = doc.add_paragraph()
                    p.add_run(f"Type: ").bold = True
                    p.add_run(f"{supplier['type']}\n")
                    p.add_run(f"Region: ").bold = True
                    p.add_run(f"{supplier['region']}\n")
                    p.add_run(f"Quality Rating: ").bold = True
                    p.add_run(f"{supplier['quality_rating']}/5\n")
                    p.add_run(f"Estimated Cost: ").bold = True
                    p.add_run(f"{currency_symbol}{supplier['estimated_unit_cost']:,.2f}/unit\n")
                    p.add_run(f"MOQ: ").bold = True
                    p.add_run(f"{supplier['moq_range'][0]:,}-{supplier['moq_range'][1]:,} units\n")
        
        # Module 5: Pricing
        if 'module_5_pricing' in results:
            doc.add_heading('Pricing & Profitability', level=1)
            pricing = results['module_5_pricing']
            amazon_pricing = pricing.get('amazon_pricing', {})
            currency_symbol = input_data.get('currency_symbol', '$')
            
            p = doc.add_paragraph()
            p.add_run(f"Product Cost: ").bold = True
            p.add_run(f"{currency_symbol}{pricing.get('product_cost', 0):,.2f}\n")
            p.add_run(f"Recommended Amazon Price: ").bold = True
            p.add_run(f"{currency_symbol}{amazon_pricing.get('recommended_price', 0):,.2f}\n")
            p.add_run(f"Profit Margin: ").bold = True
            p.add_run(f"{amazon_pricing.get('profit_margin', 0):.1f}%\n")
            p.add_run(f"ROI: ").bold = True
            p.add_run(f"{amazon_pricing.get('roi', 0):.1f}%\n")
        
        # Module 6: Platforms
        if 'module_6_platforms' in results:
            doc.add_heading('Platform Recommendations', level=1)
            platforms = results['module_6_platforms'].get('recommended_platforms', [])
            
            for i, platform in enumerate(platforms[:3], 1):
                doc.add_heading(f"{i}. {platform['platform_name']}", level=2)
                p = doc.add_paragraph()
                p.add_run(f"Recommendation Score: ").bold = True
                p.add_run(f"{platform['recommendation_score']}/100\n")
                p.add_run(f"Suitability: ").bold = True
                p.add_run(f"{platform['suitability_rating']}\n")
                
                doc.add_paragraph("Pros:", style='List Bullet')
                for pro in platform.get('pros', [])[:3]:
                    doc.add_paragraph(pro, style='List Bullet 2')
        
        # Module 7: Audience
        if 'module_7_audience' in results:
            doc.add_heading('Target Audience Profile', level=1)
            audience = results['module_7_audience'].get('audience_profile', {})
            demographics = audience.get('primary_demographics', {})
            
            p = doc.add_paragraph()
            p.add_run(f"Age Range: ").bold = True
            p.add_run(f"{demographics.get('age_range', 'N/A')}\n")
            p.add_run(f"Gender: ").bold = True
            p.add_run(f"{demographics.get('gender', 'N/A')}\n")
            p.add_run(f"Income Level: ").bold = True
            p.add_run(f"{demographics.get('income_level', 'N/A')}\n")
            
            personas = audience.get('buyer_personas', [])
            if personas:
                doc.add_heading('Primary Buyer Persona:', level=2)
                persona = personas[0]
                p = doc.add_paragraph()
                p.add_run(f"Name: ").bold = True
                p.add_run(f"{persona['name']}\n")
                p.add_run(f"Occupation: ").bold = True
                p.add_run(f"{persona['occupation']}\n")
                p.add_run(f"Goals: ").bold = True
                p.add_run(f"{', '.join(persona.get('goals', [])[:3])}\n")
        
        # Module 8: Marketing
        if 'module_8_marketing' in results:
            doc.add_heading('Marketing Strategy', level=1)
            marketing = results['module_8_marketing'].get('marketing_strategy', {})
            overview = marketing.get('overview', {})
            currency_symbol = input_data.get('currency_symbol', '$')
            
            p = doc.add_paragraph()
            p.add_run(f"Total Budget: ").bold = True
            p.add_run(f"{currency_symbol}{overview.get('total_budget', 0):,.0f}\n")
            p.add_run(f"Budget Tier: ").bold = True
            p.add_run(f"{overview.get('budget_tier', 'N/A')}\n")
            p.add_run(f"Strategy Period: ").bold = True
            p.add_run(f"{overview.get('strategy_period', 'N/A')}\n")
            
            channels = marketing.get('channel_strategy', {}).get('channels', [])
            if channels:
                doc.add_heading('Marketing Channels:', level=2)
                for channel in channels[:5]:
                    doc.add_paragraph(
                        f"{channel['channel']} - {channel['priority']} Priority ({channel['budget_allocation']})",
                        style='List Bullet'
                    )
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.add_run('Generated by E-Commerce Business Automation Platform').italic = True
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save document
        filepath.parent.mkdir(parents=True, exist_ok=True)
        doc.save(filepath)
        return True
        
    except Exception as e:
        print(f"Error exporting to DOCX: {e}")
        return False
